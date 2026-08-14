import hashlib
from pathlib import Path

from docx import Document
from PIL import Image

from pii_redactor.docx_io import DocxPackage, ImageReference
from pii_redactor.image_redactor import ImageRedactor, _Inspection


def test_all_image_policy_replaces_media_without_removing_drawing(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "redacted.docx"
    picture = tmp_path / "logo.png"
    Image.new("RGB", (180, 80), "red").save(picture)
    document = Document()
    document.add_paragraph("Issuer Limited")
    document.add_picture(str(picture))
    document.save(source)

    package = DocxPackage(source)
    references = package.image_references()
    assert len(references) == 1
    media_path = references[0].media_path
    original_hash = hashlib.sha256(package.entries[media_path]).hexdigest()
    redactor = ImageRedactor("all")
    redactions = redactor.inspect(package, [])
    assert len(redactions) == 1
    redactor.apply(package, redactions)
    package.write(output)

    reopened = DocxPackage(output)
    assert len(reopened.image_references()) == 1
    assert hashlib.sha256(reopened.entries[media_path]).hexdigest() != original_hash


def test_identity_document_ocr_markers_are_sensitive() -> None:
    assert ImageRedactor._is_identity_document(
        "INCOME TAX DEPARTMENT Permanent Account Number Card ABCDE1234F"
    )
    assert ImageRedactor._is_identity_document(
        "Unique Identification Authority of India 2943 6593 3461"
    )
    assert not ImageRedactor._is_identity_document("Quarterly sales chart")


def test_party_context_and_similar_logo_propagation() -> None:
    reference = ImageReference(
        "word/document.xml", "word/media/logo1.png", 1,
        "KSH INTERNATIONAL LIMITED CORPORATE IDENTITY NUMBER", "logo",
    )
    first = _Inspection("word/media/logo1.png", b"first", [reference], perceptual_hash=0x1234)
    second = _Inspection(
        "word/media/logo2.png", b"second",
        [ImageReference("word/document.xml", "word/media/logo2.png", 2, "", "logo")],
        perceptual_hash=0x1235,
    )
    ImageRedactor._match_sensitive_text(
        first, {"KSH International Limited"}, set()
    )
    ImageRedactor._propagate_similar_party_logos([first, second])
    assert any(reason.startswith("party-context:") for reason in first.reasons)
    assert "similar-to-redacted-party-logo" in second.reasons
