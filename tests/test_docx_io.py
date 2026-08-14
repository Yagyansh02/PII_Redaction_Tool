from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
import pytest

from pii_redactor.docx_io import DocxPackage, W_INSTR


def make_docx(path: Path) -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("alice")
    paragraph.add_run("@example.com")
    field_paragraph = document.add_paragraph("Link: ")
    instruction = OxmlElement("w:instrText")
    instruction.text = ' HYPERLINK "mailto:alice@example.com" '
    field_paragraph._p.append(instruction)
    document.save(path)


def test_cross_run_and_field_code_replacement(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    make_docx(source)
    package = DocxPackage(source)
    visible = next(record for record in package.records if record.text == "alice@example.com")
    field = next(record for record in package.records if "mailto:" in record.text)
    assert any(node.node.tag == W_INSTR for node in field.nodes)
    DocxPackage.apply_spans(visible, [(0, len(visible.text), "mira@example.com")])
    start = field.text.index("alice@example.com")
    DocxPackage.apply_spans(field, [(start, start + len("alice@example.com"), "mira@example.com")])
    package.mark_redacted()
    package.write(output)
    reopened = DocxPackage(output)
    joined = "\n".join(record.text for record in reopened.records)
    assert "alice@example.com" not in joined
    assert joined.count("mira@example.com") == 2
    assert reopened.is_redacted()


def test_xml_space_is_preserved(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    make_docx(source)
    package = DocxPackage(source)
    record = next(record for record in package.records if record.text == "alice@example.com")
    DocxPackage.apply_spans(record, [(0, 5, " Mira ")])
    assert record.nodes[0].node.get("{http://www.w3.org/XML/1998/namespace}space") == "preserve"


def test_tabs_are_virtual_hard_boundaries(tmp_path: Path) -> None:
    source = tmp_path / "tab.docx"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("2.")
    tab = OxmlElement("w:tab")
    paragraph.runs[0]._r.append(tab)
    paragraph.add_run("Compensation Committee")
    document.save(source)
    package = DocxPackage(source)
    record = next(item for item in package.records if "Compensation" in item.text)
    assert record.text == "2.\tCompensation Committee"
    assert record.crosses_hard_boundary(0, len(record.text))
    with pytest.raises(ValueError, match="crosses a tab/line break"):
        DocxPackage.apply_spans(record, [(0, len(record.text), "unsafe")])
