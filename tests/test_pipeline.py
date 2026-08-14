from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK

from evaluation.build_fixture import build_fixture
from pii_redactor.config import RedactionConfig
from pii_redactor.docx_io import DocxPackage
from pii_redactor.pipeline import RedactionPipeline


def test_synthetic_end_to_end_and_idempotency(tmp_path: Path) -> None:
    source = tmp_path / "synthetic.docx"
    gold = tmp_path / "gold.jsonl"
    corpus = tmp_path / "corpus.jsonl"
    output = tmp_path / "redacted.docx"
    mapping = tmp_path / "map.json"
    detections = tmp_path / "detections.jsonl"
    build_fixture(source, gold, corpus)
    pipeline = RedactionPipeline(RedactionConfig(use_ner=False))
    result = pipeline.run(source, output, mapping, detections)
    for expected_type in {
        "PERSON", "COMPANY", "EMAIL", "PHONE", "POSTAL_ADDRESS", "DATE_OF_BIRTH",
        "SSN", "CREDIT_CARD", "IP_ADDRESS", "PAN",
    }:
        assert result.counts.get(expected_type, 0) >= 1
    redacted_text = "\n".join(record.text for record in DocxPackage(output).records)
    for secret in ("Priya Sharma", "priya.sharma@orchid.example.org", "123-45-6789", "192.168.10.24"):
        assert secret not in redacted_text
    assert "6350960932" in redacted_text
    assert "1094535838831" in redacted_text
    second = pipeline.run(output, tmp_path / "second.docx")
    assert second.already_redacted
    assert second.total == 0


def test_pipeline_skips_tab_url_and_projects_break_split_company(tmp_path: Path) -> None:
    source = tmp_path / "boundaries.docx"
    output = tmp_path / "redacted.docx"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("2.").add_tab()
    paragraph.add_run("Compensation Committee www.example.").add_tab()
    paragraph.add_run("com")
    company = document.add_paragraph()
    company.add_run("Director, Automatic Data").add_break(WD_BREAK.LINE)
    company.add_run("Processing, Inc.")
    document.save(source)

    pipeline = RedactionPipeline(
        RedactionConfig(use_ner=False, company_scope="all", image_policy="none")
    )
    result = pipeline.run(source, output)
    assert result.skipped_hard_boundaries == 1
    text = "\n".join(record.text for record in DocxPackage(output).records)
    assert "2.\tCompensation Committee www.example.\tcom" in text
    assert "Automatic Data" not in text
    assert "Processing, Inc." not in text
